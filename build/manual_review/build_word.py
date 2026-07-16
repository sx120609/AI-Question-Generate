from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
INPUT = ROOT / "outputs" / "l2_questions.tsv"
OUT_DIR = ROOT / "outputs" / "manual_review"
OUT_PATH = OUT_DIR / "L2题目人工审阅稿.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "111827"
MUTED = "4B5563"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
BORDER = "D1D5DB"


def dxa(inches: float) -> int:
    return int(round(inches * 1440))


def set_run_font(run, name="Calibri", size=11, color=INK, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def set_style(style, size=11, color=INK, bold=False, before=0, after=6, line=1.25):
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent))

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_label_para(doc, label, text, style="Normal"):
    p = doc.add_paragraph(style=style)
    r = p.add_run(label + "：")
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_body_para(doc, text, color=INK):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, color=color)
    return p


def clean_step(step):
    return re.sub(r"^\s*\d+[）).]\s*", "", step).strip()


def split_steps(text):
    parts = [clean_step(x) for x in re.split(r"；(?=\d+[）).])", text) if x.strip()]
    return parts


def parse_tsv():
    with INPUT.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f, delimiter="\t")
        return list(reader)


def short_question(text):
    if "睡眠打卡 App" in text or "睡眠健康" in text:
        return "睡眠健康 App 上架前数据安全与素材合规预审"
    return "科罗拉多酒店停车场 EV 充电桩选址与税务预审"


def question_key(text):
    return "睡眠健康App上架预审" if "睡眠打卡 App" in text or "睡眠健康" in text else "酒店EV充电桩预审"


def urls(text):
    return re.findall(r"https?://[^，；\s]+", text)


def split_attachments(row):
    attachment_files = {
        "睡眠健康App上架预审": [
            "Apple_App_Store审核指南_应用安全隐私与法律要求.html",
            "Apple_App_Store产品页元数据创建说明.html",
            "Apple_App隐私详情与数据类型说明.html",
            "Google_Play数据安全表单填写说明.html",
            "Google_Play健康应用类别与声明要求.html",
            "FTC_背书与推荐广告指南_2023.pdf",
        ],
        "酒店EV充电桩预审": [
            "AFDC_替代燃料站数据下载说明.html",
            "AFDC_替代燃料站API说明_All_Stations.html",
            "AFDC_科罗拉多EV充电站数据_2026-07-08.csv",
            "Joint_Office_公共EV充电站选址检查清单.pdf",
            "IRS_Form8911说明_替代燃料车辆加注设施抵免_2025版.pdf",
            "IRS_企业替代燃料车辆加注设施抵免说明_2026-06.pdf",
        ],
    }
    folder_by_question = {
        "睡眠健康App上架预审": "睡眠健康App上架前数据安全与素材合规预审",
        "酒店EV充电桩预审": "科罗拉多酒店停车场EV充电桩选址与税务预审",
    }
    key = question_key(row["题目"])
    names = [x for x in row["相关附件"].split("；") if x]
    contents = [x for x in row["附件内容"].split("；") if x]
    result = []
    for i, name in enumerate(names):
        content = contents[i] if i < len(contents) else ""
        local_path = ROOT / "outputs" / "attachments" / folder_by_question[key] / attachment_files[key][i]
        result.append((i + 1, name, content, urls(content), str(local_path)))
    return result


def add_overview_table(doc, rows):
    doc.add_heading("总览", level=1)
    table = doc.add_table(rows=1, cols=5)
    headers = ["序号", "题目", "一级目录", "人类时间", "放行状态"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        shade_cell(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run_font(r, bold=True, color=DARK_BLUE)
    for idx, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        values = [str(idx), short_question(row["题目"]), row["一级目录"], row["人类完成时间"], "放行"]
        for i, value in enumerate(values):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    set_run_font(r)
    set_table_geometry(table, [700, 3200, 2200, 1400, 1500])
    set_table_borders(table)


def add_gate_table(doc):
    doc.add_heading("最终放行校验", level=1)
    rows = [
        ["校验项", "睡眠健康App上架预审", "酒店EV充电桩预审"],
        ["飞书列完整性", "通过：15列，UID留空", "通过：15列，UID留空"],
        ["L2难度", "通过：6个附件，12h，10步", "通过：6个附件，14h，11步"],
        ["主决策", "是否进入提审前改稿", "是否进入承包商现场踏勘"],
        ["附件质量", "Apple、Google Play、FTC 官方资料组合", "AFDC CSV、Joint Office 和 IRS 资料组合"],
        ["产物真实性", "Word/飞书文档 + Excel/飞书表格 + 改稿清单", "Excel/飞书表格 + Word/飞书文档 + 踏勘邮件"],
        ["证据边界", "医生背书、临床效果、审核结论为需确认", "报价、电价、抵免金额、实时状态为待确认"],
        ["最终结论", "放行", "放行"],
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    for r_idx, values in enumerate(rows):
        for c_idx, value in enumerate(values):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = value
            shade_cell(cell, LIGHT_BLUE if r_idx == 0 else "FFFFFF")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, bold=(r_idx == 0), color=DARK_BLUE if r_idx == 0 else INK)
    set_table_geometry(table, [2000, 3650, 3710])
    set_table_borders(table)


def add_question_section(doc, idx, row):
    title = short_question(row["题目"])
    doc.add_heading(f"题目 {idx}：{title}", level=1)
    add_label_para(doc, "一级目录", row["一级目录"])
    add_label_para(doc, "二级/三级目录", f"{row['二级目录']} / {row['三级目录']}")
    add_label_para(doc, "任务概括", row["任务概括"])
    add_label_para(doc, "专家年限与人类时间", f"{row['标注专家工作年限']}；{row['人类完成时间']}")

    doc.add_heading("题面", level=2)
    add_body_para(doc, row["题目"])

    doc.add_heading("交付物检查", level=2)
    add_label_para(doc, "产物格式", row["产物格式"])
    add_label_para(doc, "产物内容", row["产物内容"])

    doc.add_heading("做题关键步骤", level=2)
    for n, step in enumerate(split_steps(row["做题关键步骤"]), start=1):
        p = doc.add_paragraph(style="Normal")
        r = p.add_run(f"{n}. ")
        set_run_font(r, bold=True, color=DARK_BLUE)
        r = p.add_run(step)
        set_run_font(r)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)

    doc.add_heading("附件来源与边界", level=2)
    for seq, name, content, found_urls, local_path in split_attachments(row):
        doc.add_heading(f"资料 {seq}：{name}", level=3)
        add_label_para(doc, "本地文件", local_path)
        if found_urls:
            add_label_para(doc, "URL", "\n".join(found_urls))
        add_label_para(doc, "用途与边界", content)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    rows = parse_tsv()
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    set_style(doc.styles["Normal"], size=11, color=INK, before=0, after=6, line=1.25)
    set_style(doc.styles["Heading 1"], size=16, color=BLUE, bold=True, before=18, after=10, line=1.25)
    set_style(doc.styles["Heading 2"], size=13, color=BLUE, bold=True, before=14, after=7, line=1.25)
    set_style(doc.styles["Heading 3"], size=12, color=DARK_BLUE, bold=True, before=10, after=5, line=1.25)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("L2题目人工审阅稿")
    set_run_font(run, size=9, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("L2题目人工审阅稿")
    set_run_font(run, size=22, color="0B2545", bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("生成日期：2026-07-09 | 题目数：2 | 用途：人工查看与复核")
    set_run_font(run, size=10, color=MUTED)

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(10)
    run = note.add_run("说明：本文件是内部审阅稿。飞书可粘贴字段以 Excel/TSV 为准；题目要求模型交付的产物没有使用 Markdown，均改为真实职场交付形式。")
    set_run_font(run, size=10, color=INK, bold=True)

    add_overview_table(doc, rows)
    add_gate_table(doc)
    for idx, row in enumerate(rows, start=1):
        add_question_section(doc, idx, row)
        if idx < len(rows):
            doc.add_section(WD_SECTION.NEW_PAGE)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
